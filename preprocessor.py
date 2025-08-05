import fitz  # PDF
import docx  # DOCX
import pandas as pd  # CSV
import sqlite3  # SQLite databases
import pytesseract  # OCR
from PIL import Image  # Image processing
import io
import requests
from pathlib import Path
from typing import List, Dict, Union, Optional
from langchain_core.documents import Document
from langchain_text_splitters import CharacterTextSplitter

class FilePreprocessor:
    def __init__(self):
        self.supported_extensions = {
            '.pdf', '.txt', '.docx', '.csv', 
            '.db', '.sqlite', '.sqlite3',
            '.jpg', '.jpeg', '.png'
        }
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Update this path as needed

    def process_file(self, file_input: Union[str, Path], 
                    is_url: bool = False,
                    chunk_size: int = 1000,
                    chunk_overlap: int = 200,
                    metadata: Optional[Dict] = None) -> List[Document]:
        """Process files into LangChain Documents with metadata."""
        content = self._extract_content(file_input, is_url)
        chunks = self._chunk_content(content, chunk_size, chunk_overlap)
        
        return [
            Document(
                page_content=chunk,
                metadata=metadata or {"source": str(file_input)}
            )
            for chunk in chunks
        ]

    def _extract_content(self, file_input: Union[str, Path], 
                         is_url: bool) -> str:
        """Extract content from various file formats."""
        if is_url:
            return self._extract_from_url(file_input)
        
        file_path = Path(file_input)
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        if file_ext == '.pdf':
            text = self._extract_from_pdf(file_path)
            if not text.strip():
                text = self._extract_with_ocr(file_path)
            return text
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext == '.docx':
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_ext == '.csv':
            return pd.read_csv(file_path).to_string()
        elif file_ext in ('.db', '.sqlite', '.sqlite3'):
            return self._extract_from_db(file_path)
        elif file_ext in ('.jpg', '.jpeg', '.png'):
            return self._extract_with_ocr(file_path)

    def _extract_from_url(self, url: str) -> str:
        """Extract content from URL."""
        response = requests.get(url)
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', '').lower()
        
        if 'pdf' in content_type:
            with io.BytesIO(response.content) as pdf_file:
                doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
                text = " ".join(page.get_text() for page in doc)
                if not text.strip():
                    text = self._extract_with_ocr(pdf_file)
                return text
        elif any(img_type in content_type for img_type in ['jpg', 'jpeg', 'png']):
            image = Image.open(io.BytesIO(response.content))
            return pytesseract.image_to_string(image)
        else:
            return response.text

    def _extract_from_pdf(self, pdf_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        doc = fitz.open(pdf_path)
        return " ".join(page.get_text() for page in doc)

    def _extract_with_ocr(self, file_input: Union[Path, io.BytesIO]) -> str:
        """Extract text using OCR."""
        if isinstance(file_input, (Path, str)):
            image = Image.open(file_input)
        else:  # BytesIO
            image = Image.open(file_input)
        return pytesseract.image_to_string(image)

    def _extract_from_db(self, db_path: Path) -> str:
        """Extract schema and sample data from SQLite database."""
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
        
        result = []
        for table in tables:
            table_name = table[0]
            result.append(f"\nTable: {table_name}")
            
            cursor.execute(f"PRAGMA table_info({table_name});")
            columns = cursor.fetchall()
            result.append("Columns:")
            for col in columns:
                result.append(f"  {col[1]} ({col[2]})")
            
            cursor.execute(f"SELECT * FROM {table_name} LIMIT 5;")
            rows = cursor.fetchall()
            if rows:
                result.append("Sample data:")
                for row in rows:
                    result.append(f"  {row}")
        
        conn.close()
        return "\n".join(result)

    def _chunk_content(self, content: str, 
                       chunk_size: int, 
                       chunk_overlap: int) -> List[str]:
        """Split content into chunks."""
        text_splitter = CharacterTextSplitter(
            separator="\n\n",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        return text_splitter.split_text(content)